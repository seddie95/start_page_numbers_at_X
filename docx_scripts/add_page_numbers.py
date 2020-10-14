import time
from docx import Document
from docx.oxml import OxmlElement, ns
import os
from specific_page import settings


def create_attribute(element, name, value):
    """ Create attribute for a given element """
    element.set(ns.qn(name), value)


def set_page_number_type(fmt='', start_num='1'):
    """ specify the starting page number and style"""
    # Add numbering to section
    num_type = OxmlElement('w:pgNumType')
    create_attribute(num_type, 'w:start', start_num)

    # Set the number format
    if fmt != '':
        create_attribute(num_type, 'w:fmt', fmt)

    return num_type


def add_section(parag, fmt=''):
    """ Create new section tag element"""
    pPr = OxmlElement('w:pPr')
    section = OxmlElement('w:sectPr')

    # set the page number type
    num_type = set_page_number_type(fmt)

    # Append the section to the paragraph
    section.append(num_type)
    pPr.append(section)
    parag._p.append(pPr)


def add_page_number(parag, position=''):
    """ Add page numbers to the footer of the document"""
    run = parag.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = OxmlElement('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    # Set the position for the numbers
    if position != '':
        jc = OxmlElement('w:jc')
        create_attribute(jc, 'w:val', position)
        p = parag._p
        pPr = p.get_or_add_pPr()
        pPr.append(jc)

    # Append the new create elements to the r tag
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_page_size(sect1, sect2):
    """ Ensure that both sections have the same size"""
    sect1.page_height = sect2.page_height
    sect1.page_width = sect2.page_width
    sect1.left_margin = sect2.left_margin
    sect1.right_margin = sect2.right_margin
    sect1.top_margin = sect2.top_margin
    sect1.bottom_margin = sect2.bottom_margin
    sect1.header_distance = sect2.header_distance
    sect1.footer_distance = sect2.footer_distance


def save_file(doc, doc_obj, file_name):
    """ Save a copy of the file with the page numbers """
    saved_file = f'{settings.MEDIA_ROOT}/numbered_{file_name}'
    doc.save(saved_file)

    # Delete object from database and server
    doc_obj.delete()

    while not os.path.exists(saved_file):
        time.sleep(1)

    if os.path.isfile(saved_file):
        return f'media/numbered_{file_name}'

    else:
        return 'File does not exist!'


def set_page_numbers(specs):
    doc_obj = specs['doc_obj']
    file_path = doc_obj.doc_file.path
    file_name = doc_obj.file_name
    paragraph_number = int(specs['heading'][0])
    position = specs['position'][0]
    style = specs['style'][0]

    # Open existing and find last paragraph before section break
    try:
        if os.path.exists(file_path):
            doc = Document(file_path)
            # Set the page number type so that it starts from one

            # Add normal page number styling if paragraph is zero
            if paragraph_number == 0:
                add_page_number(doc.sections[0].footer.paragraphs[0], position)

                # Save the file and delete unneeded file
                return save_file(doc, doc_obj, file_name)

            else:
                sect = doc.sections[0]._sectPr
                pg_num_type = set_page_number_type()
                sect.append(pg_num_type)

                # add numbers starting at i
                chosen_paragraph = doc.paragraphs[paragraph_number]
                new_paragraph = chosen_paragraph.insert_paragraph_before()
                add_section(new_paragraph, style)

                # set the page to be A4
                set_page_size(doc.sections[0], doc.sections[1])

                # Add the page numbers
                add_page_number(doc.sections[0].footer.paragraphs[0], position)

                # Save the file and delete unneeded file
                return save_file(doc, doc_obj, file_name)

    except OSError:
        return 1
